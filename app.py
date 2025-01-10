from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from PyPDF2 import PdfMerger
from PIL import Image
from docx import Document
import io
import csv

app = Flask(__name__)
CORS(app)  # Enable CORS for all origins

ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'jpg', 'jpeg', 'png', 'csv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/api/merge', methods=['POST'])
def merge_files():
    try:
        # Validate input
        if 'files' not in request.files or 'file_type' not in request.form:
            return jsonify({'error': 'No files or file type provided'}), 400

        file_type = request.form['file_type']
        files = request.files.getlist('files')

        if len(files) < 2:
            return jsonify({'error': 'Please upload at least 2 files.'}), 400

        # Check file extensions
        if file_type not in ALLOWED_EXTENSIONS:
            return jsonify({'error': 'Unsupported file type.'}), 400

        if not all(allowed_file(file.filename) for file in files):
            return jsonify({'error': 'File type mismatch detected.'}), 400

        # Merge files based on type
        if file_type == 'pdf':
            return merge_pdfs(files)
        elif file_type in {'jpeg', 'png'}:
            return merge_images(files)
        elif file_type == 'txt':
            return merge_texts(files)
        elif file_type == 'csv':
            return merge_csvs(files)
        elif file_type == 'doc':
            return merge_word_docs(files)
        else:
            return jsonify({'error': 'Unsupported file type.'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def merge_images(files):
    try:
        # Determine the image format from the first file
        first_file_extension = files[0].filename.rsplit('.', 1)[1].lower()
        if first_file_extension not in {'jpeg', 'jpg', 'png'}:
            raise Exception("Invalid image format for merging.")

        images = [Image.open(file.stream).convert("RGB") for file in files]
        
        # Output as a single PDF
        output = io.BytesIO()
        images[0].save(output, format="PDF", save_all=True, append_images=images[1:])
        output.seek(0)

        return send_file(output, mimetype='application/pdf', as_attachment=True, download_name='merged_images.pdf')
    except Exception as e:
        raise Exception(f"Image merging failed: {e}")

def merge_pdfs(files):
    try:
        merger = PdfMerger()
        for file in files:
            merger.append(file.stream)
        output = io.BytesIO()
        merger.write(output)
        output.seek(0)
        merger.close()
        return send_file(output, mimetype='application/pdf', as_attachment=True, download_name='merged.pdf')
    except Exception as e:
        return jsonify({'error': f'PDF merging failed: {str(e)}'}), 500

def merge_texts(files):
    try:
        merged_content = ""
        for file in files:
            merged_content += file.read().decode('utf-8') + "\n"
        output = io.BytesIO(merged_content.encode('utf-8'))
        return send_file(output, mimetype='text/plain', as_attachment=True, download_name='merged.txt')
    except Exception as e:
        return jsonify({'error': f'Text merging failed: {str(e)}'}), 500

def merge_csvs(files):
    try:
        merged_rows = []
        for file in files:
            reader = csv.reader(io.StringIO(file.read().decode('utf-8')))
            merged_rows.extend(reader)
        output = io.BytesIO()
        writer = csv.writer(io.TextIOWrapper(output, write_through=True))
        writer.writerows(merged_rows)
        output.seek(0)
        return send_file(output, mimetype='text/csv', as_attachment=True, download_name='merged.csv')
    except Exception as e:
        return jsonify({'error': f'CSV merging failed: {str(e)}'}), 500

def merge_word_docs(files):
    try:
        merged_doc = Document()
        for file in files:
            doc = Document(file.stream)
            for paragraph in doc.paragraphs:
                merged_doc.add_paragraph(paragraph.text)
        output = io.BytesIO()
        merged_doc.save(output)
        output.seek(0)
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name='merged.docx')
    except Exception as e:
        return jsonify({'error': f'Word document merging failed: {str(e)}'}), 500

@app.route('/api/files/metadata', methods=['POST'])
def file_metadata():
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'No files uploaded'}), 400

        files = request.files.getlist('files')
        metadata = []
        for file in files:
            metadata.append({
                'name': file.filename,
                'size': len(file.read()),  # File size in bytes
            })

        return jsonify({'metadata': metadata}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
